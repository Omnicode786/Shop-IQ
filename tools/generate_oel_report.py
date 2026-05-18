from __future__ import annotations

import math
import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "generated_reports"
ASSET_DIR = OUT / "assets"
DOCX = OUT / "ShopIQ_OEL_DBMS_Report.docx"

DEEP = "121624"
INK = "1F2937"
MUTED = "64748B"
TEAL = "10B981"
CYAN = "06B6D4"
AMBER = "F59E0B"
RED = "EF4444"
BLUE = "3B82F6"
SURFACE = "F8FAFC"
LINE = "CBD5E1"


def ensure_dirs():
    OUT.mkdir(exist_ok=True)
    ASSET_DIR.mkdir(exist_ok=True)


def font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def rgb(hex_color: str):
    hex_color = hex_color.strip("#")
    return tuple(int(hex_color[i : i + 2], 16) for i in (0, 2, 4))


def draw_rounded(draw: ImageDraw.ImageDraw, xy, fill, outline=None, radius=24, width=2):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def center_text(draw: ImageDraw.ImageDraw, box, text, fnt, fill):
    x1, y1, x2, y2 = box
    bbox = draw.textbbox((0, 0), text, font=fnt)
    x = x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2
    y = y1 + (y2 - y1 - (bbox[3] - bbox[1])) / 2 - 2
    draw.text((x, y), text, font=fnt, fill=fill)


def arrow(draw, start, end, fill, width=5):
    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 16
    p1 = (end[0] - size * math.cos(angle - math.pi / 7), end[1] - size * math.sin(angle - math.pi / 7))
    p2 = (end[0] - size * math.cos(angle + math.pi / 7), end[1] - size * math.sin(angle + math.pi / 7))
    draw.polygon([end, p1, p2], fill=fill)


def architecture_image():
    path = ASSET_DIR / "shopiq_architecture.png"
    img = Image.new("RGB", (1500, 860), rgb("F8FAFC"))
    draw = ImageDraw.Draw(img)
    title_f = font(44, True)
    label_f = font(28, True)
    small_f = font(22)
    tiny_f = font(18)

    draw.text((70, 50), "ShopIQ System Architecture", font=title_f, fill=rgb(DEEP))
    draw.text((72, 104), "A role-aware Next.js workspace connected to PostgreSQL through Prisma and secured API routes.", font=small_f, fill=rgb(MUTED))

    boxes = [
        ("Users", "Admin, Manager, Staff", (80, 245, 300, 395), TEAL),
        ("Next.js UI", "Dashboards, CRUD pages, AI tab", (390, 210, 650, 430), BLUE),
        ("API Routes", "Validation, RBAC, errors", (735, 210, 995, 430), CYAN),
        ("Prisma ORM", "Typed database calls", (1080, 245, 1300, 395), AMBER),
        ("PostgreSQL", "Normalized business data", (1080, 520, 1300, 670), DEEP),
        ("Gemini Agent", "Tools + approval gates", (390, 540, 650, 700), "8B5CF6"),
    ]
    for title, subtitle, box, color in boxes:
        shadow = (box[0] + 8, box[1] + 10, box[2] + 8, box[3] + 10)
        draw_rounded(draw, shadow, rgb("E2E8F0"), radius=30, outline=None, width=0)
        draw_rounded(draw, box, "white", radius=30, outline=rgb(color), width=4)
        draw.ellipse((box[0] + 24, box[1] + 28, box[0] + 78, box[1] + 82), fill=rgb(color))
        draw.text((box[0] + 96, box[1] + 32), title, font=label_f, fill=rgb(DEEP))
        draw.text((box[0] + 32, box[1] + 102), subtitle, font=small_f, fill=rgb(MUTED))

    arrow(draw, (302, 320), (385, 320), rgb(LINE), 6)
    arrow(draw, (652, 320), (730, 320), rgb(LINE), 6)
    arrow(draw, (997, 320), (1076, 320), rgb(LINE), 6)
    arrow(draw, (1190, 398), (1190, 516), rgb(LINE), 6)
    arrow(draw, (520, 536), (520, 434), rgb("8B5CF6"), 5)
    arrow(draw, (650, 620), (735, 390), rgb("8B5CF6"), 5)
    arrow(draw, (995, 390), (1080, 560), rgb(CYAN), 5)

    draw_rounded(draw, (80, 735, 1300, 805), rgb("ECFDF5"), radius=24, outline=rgb("BBF7D0"), width=2)
    draw.text((115, 756), "Core guarantees: authenticated sessions, role checks on every route, Zod validation, normalized relations, and preview-first AI writes.", font=tiny_f, fill=rgb("166534"))
    img.save(path)
    return path


def erd_image():
    path = ASSET_DIR / "shopiq_erd.png"
    img = Image.new("RGB", (1500, 1050), rgb("FBFCFF"))
    draw = ImageDraw.Draw(img)
    title_f = font(42, True)
    label_f = font(24, True)
    body_f = font(18)
    draw.text((70, 44), "Simplified Entity Relationship View", font=title_f, fill=rgb(DEEP))
    draw.text((72, 94), "Main relational structure used by ShopIQ. Operational rows are scoped by shopId for multi-shop safety.", font=body_f, fill=rgb(MUTED))

    entity_boxes = {
        "Shop": (625, 165, 875, 285, TEAL, ["id PK", "name, city, currency"]),
        "User": (165, 175, 430, 310, BLUE, ["id PK", "shopId FK", "role, status"]),
        "Category": (1040, 175, 1310, 310, CYAN, ["id PK", "shopId FK", "unique(shopId,name)"]),
        "Product": (960, 400, 1355, 575, AMBER, ["id PK", "shopId/categoryId FK", "sku unique per shop", "prices, stock, reorder"]),
        "Customer": (115, 420, 455, 575, "8B5CF6", ["id PK", "shopId FK", "creditLimit, balance"]),
        "Supplier": (115, 680, 455, 835, RED, ["id PK", "shopId FK", "balance, reliabilityScore"]),
        "Invoice": (555, 410, 820, 555, BLUE, ["id PK", "customerId FK", "createdById FK", "status, totals"]),
        "InvoiceItem": (555, 660, 820, 805, CYAN, ["id PK", "invoiceId FK", "productId FK", "quantity, total"]),
        "Purchase": (960, 670, 1225, 815, TEAL, ["id PK", "supplierId FK", "createdById FK", "status, totals"]),
        "PurchaseItem": (960, 875, 1225, 1000, AMBER, ["id PK", "purchaseId FK", "productId FK"]),
        "Payment": (505, 875, 850, 1010, "0EA5E9", ["customer/supplier", "invoice/purchase", "direction, amount"]),
        "StockMovement": (1225, 680, 1440, 835, "475569", ["productId FK", "type, before/after qty"]),
    }
    centers = {}
    for name, (x1, y1, x2, y2, color, fields) in entity_boxes.items():
        centers[name] = ((x1 + x2) // 2, (y1 + y2) // 2)
        draw_rounded(draw, (x1 + 6, y1 + 8, x2 + 6, y2 + 8), rgb("E2E8F0"), radius=22, outline=None, width=0)
        draw_rounded(draw, (x1, y1, x2, y2), "white", radius=22, outline=rgb(color), width=3)
        draw.rounded_rectangle((x1, y1, x2, y1 + 44), radius=22, fill=rgb(color))
        draw.rectangle((x1, y1 + 22, x2, y1 + 44), fill=rgb(color))
        center_text(draw, (x1, y1, x2, y1 + 44), name, label_f, "white")
        fy = y1 + 58
        for field in fields[:4]:
            draw.text((x1 + 18, fy), field, font=body_f, fill=rgb(INK))
            fy += 24

    rels = [
        ("Shop", "User"), ("Shop", "Category"), ("Category", "Product"), ("Shop", "Product"),
        ("Customer", "Invoice"), ("User", "Invoice"), ("Invoice", "InvoiceItem"), ("Product", "InvoiceItem"),
        ("Supplier", "Purchase"), ("User", "Purchase"), ("Purchase", "PurchaseItem"), ("Product", "PurchaseItem"),
        ("Product", "StockMovement"), ("Invoice", "Payment"), ("Purchase", "Payment"), ("Customer", "Payment"), ("Supplier", "Payment"),
    ]
    for a, b in rels:
        arrow(draw, centers[a], centers[b], rgb("94A3B8"), 3)

    img.save(path)
    return path


def rubric_image():
    path = ASSET_DIR / "rubric_fit.png"
    img = Image.new("RGB", (1500, 500), rgb("F8FAFC"))
    draw = ImageDraw.Draw(img)
    title_f = font(42, True)
    label_f = font(23, True)
    small_f = font(18)
    draw.text((60, 44), "Rubric Fit Snapshot", font=title_f, fill=rgb(DEEP))
    items = [
        ("Database", "14 relational models", TEAL),
        ("Keys", "PK, FK, unique, indexes", BLUE),
        ("CRUD", "module APIs + UI", CYAN),
        ("Queries", "joins, filters, reports", AMBER),
        ("Connectivity", "Next.js -> Prisma -> PostgreSQL", "8B5CF6"),
        ("Reports", "dashboard + analytics", RED),
    ]
    x = 60
    for i, (name, desc, color) in enumerate(items):
        y = 140 + (i % 2) * 150
        if i == 3:
            x = 60
        x = 60 + (i % 3) * 470
        draw_rounded(draw, (x, y, x + 420, y + 105), "white", radius=26, outline=rgb(color), width=4)
        draw.ellipse((x + 24, y + 26, x + 74, y + 76), fill=rgb(color))
        draw.text((x + 94, y + 25), name, font=label_f, fill=rgb(DEEP))
        draw.text((x + 94, y + 58), desc, font=small_f, fill=rgb(MUTED))
    img.save(path)
    return path


def ai_features_image():
    path = ASSET_DIR / "shopiq_ai_features.png"
    img = Image.new("RGB", (1500, 760), rgb("F8FAFC"))
    draw = ImageDraw.Draw(img)
    title_f = font(42, True)
    label_f = font(23, True)
    body_f = font(18)
    small_f = font(16)

    draw.text((70, 45), "ShopIQ AI Assistant Feature System", font=title_f, fill=rgb(DEEP))
    draw.text((72, 95), "Gemini is connected to role-aware ShopIQ tools. Read tasks can answer immediately; write tasks require approval.", font=body_f, fill=rgb(MUTED))

    cards = [
        ("Live business context", "Dashboard metrics, low stock, dues, revenue timeline and movement signals.", (70, 165, 455, 295), TEAL),
        ("Search and details", "Find products, customers, invoices, purchases, payments, suppliers and staff.", (555, 165, 940, 295), BLUE),
        ("Operating jobs", "Reorder plan, collections plan, cashflow risk, sales-quality review and stock audit.", (1040, 165, 1425, 295), CYAN),
        ("Exact answers", "Sales, earnings, profit, cash received, customer pending money and recent payments.", (70, 365, 455, 495), AMBER),
        ("Approval-gated writes", "Create/update records only after preview and user confirmation.", (555, 365, 940, 495), "8B5CF6"),
        ("Guardrails", "No deletes, no secrets, no permission bypass, no invented ids or balances.", (1040, 365, 1425, 495), RED),
    ]
    for title, body, box, color in cards:
        x1, y1, x2, y2 = box
        draw_rounded(draw, (x1 + 7, y1 + 9, x2 + 7, y2 + 9), rgb("E2E8F0"), radius=26, outline=None, width=0)
        draw_rounded(draw, box, "white", radius=26, outline=rgb(color), width=4)
        draw.ellipse((x1 + 26, y1 + 26, x1 + 78, y1 + 78), fill=rgb(color))
        draw.text((x1 + 96, y1 + 24), title, font=label_f, fill=rgb(DEEP))
        lines = []
        words = body.split()
        current = ""
        for word in words:
            trial = (current + " " + word).strip()
            if draw.textlength(trial, font=small_f) > (x2 - x1 - 116):
                lines.append(current)
                current = word
            else:
                current = trial
        if current:
            lines.append(current)
        for i, line in enumerate(lines[:3]):
            draw.text((x1 + 96, y1 + 60 + i * 22), line, font=small_f, fill=rgb(MUTED))

    # workflow rail
    y = 625
    steps = [
        ("Ask", TEAL),
        ("Gemini selects tool", BLUE),
        ("Role-aware data", CYAN),
        ("Preview if write", "8B5CF6"),
        ("Approve", AMBER),
        ("Prisma write", RED),
    ]
    x = 100
    prev = None
    for label, color in steps:
        box = (x, y, x + 185, y + 70)
        draw_rounded(draw, box, "white", radius=22, outline=rgb(color), width=3)
        center_text(draw, box, label, small_f, rgb(DEEP))
        if prev:
            arrow(draw, (prev[2] + 4, y + 35), (box[0] - 8, y + 35), rgb("94A3B8"), 4)
        prev = box
        x += 220

    img.save(path)
    return path


def shopiq_wordmark():
    path = ASSET_DIR / "shopiq_wordmark.png"
    img = Image.new("RGBA", (820, 260), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    title_f = font(86, True)
    sub_f = font(25)
    draw_rounded(draw, (36, 36, 194, 194), rgb("ECFDF5"), outline=rgb(TEAL), radius=44, width=5)
    draw_rounded(draw, (72, 83, 118, 158), rgb(TEAL), radius=18)
    draw_rounded(draw, (126, 61, 158, 158), rgb(CYAN), radius=16)
    draw.arc((76, 60, 168, 162), 205, 348, fill=rgb(DEEP), width=10)
    draw.ellipse((156, 55, 176, 75), fill=rgb(AMBER))
    draw.text((230, 54), "Shop", font=title_f, fill=rgb(DEEP))
    draw.text((456, 54), "IQ", font=title_f, fill=rgb(TEAL))
    draw.text((236, 154), "retail intelligence workspace", font=sub_f, fill=rgb(MUTED))
    img.save(path)
    return path


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9E2EC", size="8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_pct=100):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_pct * 50))
    tbl_w.set(qn("w:type"), "pct")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def remove_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tc_pr = tc.get_or_add_tcPr()
            borders = tc_pr.first_child_found_in("w:tcBorders")
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                element = borders.find(qn("w:" + edge))
                if element is None:
                    element = OxmlElement("w:" + edge)
                    borders.append(element)
                element.set(qn("w:val"), "nil")


def style_table(table, header_fill=DEEP, header_text="FFFFFF", band_fill="F8FAFC"):
    table.style = "Table Grid"
    set_table_width(table)
    for idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell, "D9E2EC", "6")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(9.2)
                    run.font.color.rgb = RGBColor.from_string(INK)
            if idx == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor.from_string(header_text)
            elif idx % 2 == 0:
                set_cell_shading(cell, band_fill)
    if table.rows:
        set_repeat_table_header(table.rows[0])


def p(doc, text="", style=None, align=None, color=None, size=None, bold=False):
    para = doc.add_paragraph(style=style)
    if align:
        para.alignment = align
    run = para.add_run(text)
    run.font.name = "Aptos"
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.bold = bold
    return para


def add_heading(doc, text, level=1):
    para = doc.add_paragraph()
    para.style = f"Heading {level}"
    run = para.add_run(text)
    run.font.name = "Aptos Display"
    return para


def add_callout(doc, title, body, fill="ECFDF5", accent=TEAL):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(0.32)
    table.columns[1].width = Cm(16.4)
    remove_table_borders(table)
    set_cell_shading(table.cell(0, 0), accent)
    set_cell_shading(table.cell(0, 1), fill)
    cell = table.cell(0, 1)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    title_p = cell.paragraphs[0]
    title_p.paragraph_format.space_after = Pt(2)
    r = title_p.add_run(title)
    r.bold = True
    r.font.name = "Aptos"
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(DEEP)
    body_p = cell.add_paragraph()
    body_p.paragraph_format.space_after = Pt(0)
    rr = body_p.add_run(body)
    rr.font.name = "Aptos"
    rr.font.size = Pt(9.5)
    rr.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph()


def add_kpi_strip(doc):
    data = [
        ("14", "Core models", "Shop-scoped relational schema"),
        ("3", "User roles", "Admin, Manager, Staff"),
        ("9+", "Modules", "Inventory, billing, customers, reports"),
        ("100%", "Rubric mapped", "Each OEL criterion addressed"),
    ]
    table = doc.add_table(rows=1, cols=4)
    remove_table_borders(table)
    for i, (num, label, sub) in enumerate(data):
        cell = table.cell(0, i)
        set_cell_shading(cell, ["ECFDF5", "EFF6FF", "ECFEFF", "FFF7ED"][i])
        for par in cell.paragraphs:
            par.paragraph_format.space_after = Pt(1)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cell.paragraphs[0].add_run(num)
        r.font.name = "Aptos Display"
        r.font.size = Pt(24)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string([TEAL, BLUE, CYAN, AMBER][i])
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label)
        r2.font.bold = True
        r2.font.size = Pt(9.8)
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(sub)
        r3.font.size = Pt(8.3)
        r3.font.color.rgb = RGBColor.from_string(MUTED)


def set_doc_styles(doc):
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.2)
    styles["Normal"].font.color.rgb = RGBColor.from_string(INK)
    for i, size, color in [(1, 19, DEEP), (2, 14, TEAL), (3, 11.5, BLUE)]:
        st = styles[f"Heading {i}"]
        st.font.name = "Aptos Display"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(11 if i == 1 else 8)
        st.paragraph_format.space_after = Pt(4)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(30)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = RGBColor.from_string(DEEP)


def configure_section(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)


def setup_sections(doc):
    section = doc.sections[0]
    configure_section(section)
    section.header.is_linked_to_previous = False
    section.header.paragraphs[0].text = ""
    section.footer.is_linked_to_previous = False
    section.footer.paragraphs[0].text = ""


def start_body_section(doc):
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(section)
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = "ShopIQ OEL Report | CS 222 Database Management Systems"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in hp.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(MUTED)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("ShopIQ - AI-powered inventory and sales operating system")
    for run in fp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def add_cover(doc, logo_path):
    # Minimal cover with academic identity and project branding.
    p(doc, "DEPARTMENT OF COMPUTER & INFORMATION SYSTEMS ENGINEERING", align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED, size=9, bold=True)
    p(doc, "BACHELORS IN COMPUTER SYSTEMS ENGINEERING", align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED, size=9, bold=True)
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(logo_path), width=Cm(8.4))
    p(doc, "Open Ended Lab Report", align=WD_ALIGN_PARAGRAPH.CENTER, color=TEAL, size=14, bold=True)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("ShopIQ").bold = True
    p(doc, "AI-powered inventory, billing, reporting, and staff management system", align=WD_ALIGN_PARAGRAPH.CENTER, color=INK, size=13)
    doc.add_paragraph()

    table = doc.add_table(rows=5, cols=2)
    remove_table_borders(table)
    rows = [
        ("Course", "CS 222 - Database Management Systems"),
        ("CLO / Taxonomy", "CLO-2 / C3"),
        ("Backend DBMS", "PostgreSQL using Prisma ORM"),
        ("Frontend", "Next.js 14 App Router with role-based dashboards"),
        ("Submitted Project", "ShopIQ retail operations workspace"),
    ]
    for i, (k, v) in enumerate(rows):
        set_cell_shading(table.cell(i, 0), "E2E8F0")
        set_cell_shading(table.cell(i, 1), "F8FAFC")
        table.cell(i, 0).paragraphs[0].add_run(k).bold = True
        table.cell(i, 1).paragraphs[0].add_run(v)
        for cell in table.rows[i].cells:
            set_cell_border(cell, "FFFFFF", "10")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(10.2)
                    run.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph()
    add_callout(
        doc,
        "Design position",
        "This report is structured directly around the OEL brief and evaluation rubric: database creation, keys and constraints, CRUD, query implementation, connectivity, reporting, and project understanding.",
        fill="F0FDFA",
        accent=TEAL,
    )
    p(doc, "Prepared as a final DBMS OEL submission for a working software application.", align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED, size=9)


def add_contents(doc):
    add_heading(doc, "Contents", 1)
    items = [
        "1. Project Overview",
        "2. Problem Statement and Objectives",
        "3. Technology Stack",
        "4. System Architecture",
        "5. User Roles and Permission Design",
        "6. Database Design and ERD",
        "7. Tables, Keys, Constraints and Indexes",
        "8. CRUD Operations and Validation",
        "9. Query Implementation and Reports",
        "10. Front-End and Database Connectivity",
        "11. AI Assistant Feature System",
        "12. Testing, Exception Handling and Rubric Mapping",
        "Appendix A. Representative SQL / Prisma Queries",
    ]
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(item)
    doc.add_page_break()


def add_project_overview(doc, rubric_path):
    add_heading(doc, "1. Project Overview", 1)
    p(
        doc,
        "ShopIQ is a database-backed retail operating system for shops that need one place to manage inventory, customers, suppliers, invoices, purchases, payments, reports and staff. The application is not a static prototype: it uses PostgreSQL through Prisma, authenticated Next.js route handlers, role-aware CRUD screens, and a Gemini-powered assistant that can answer business questions from live operational data.",
    )
    add_kpi_strip(doc)
    doc.add_paragraph()
    doc.add_picture(str(rubric_path), width=Cm(17.3))
    p(doc, "Figure 1. How ShopIQ maps to the OEL evaluation areas.", align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED, size=8.5)
    add_callout(
        doc,
        "Why this project fits DBMS OEL",
        "The project demonstrates schema design, integrity rules, many-to-one and one-to-many relations, CRUD operations, report queries, role separation, and a complete database-connected front end.",
        fill="EFF6FF",
        accent=BLUE,
    )
    doc.add_page_break()


def add_problem_objectives(doc):
    add_heading(doc, "2. Problem Statement and Objectives", 1)
    p(doc, "Retail shops often manage sales, stock and credit information through notebooks, spreadsheets or disconnected billing tools. That creates four recurring problems: stock is reordered late, customer dues are hard to track, supplier payments become unclear, and managers cannot quickly see the condition of the business.")
    add_heading(doc, "Project objectives", 2)
    objectives = [
        "Build a relational database system that stores shop operations in normalized tables.",
        "Provide user-friendly forms for insert, update, delete, search and reporting operations.",
        "Enforce role-based access for Admin, Manager and Staff users.",
        "Generate meaningful reports for sales, revenue, dues, stock risk and product movement.",
        "Add an AI assistant that can answer live business questions and prepare safe database actions only after approval.",
    ]
    for item in objectives:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(item)

    add_heading(doc, "Main users", 2)
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "User"
    table.rows[0].cells[1].text = "Purpose"
    table.rows[0].cells[2].text = "Typical privileges"
    rows = [
        ("Admin", "Owner-level control of the shop workspace.", "Full management of products, customers, suppliers, purchases, staff, settings and reports."),
        ("Manager", "Operational control without unrestricted ownership risk.", "Can manage most operational modules and create staff members within allowed rules."),
        ("Staff", "Daily billing, payments and customer handling.", "Can work with permitted customer, invoice, payment, inventory view and assistant features."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)
    doc.add_page_break()


def add_stack_architecture(doc, arch_path):
    add_heading(doc, "3. Technology Stack", 1)
    table = doc.add_table(rows=1, cols=4)
    headers = ["Layer", "Technology", "Role in project", "OEL evidence"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    rows = [
        ("Frontend", "Next.js 14, React, Tailwind, shadcn-style UI", "Forms, dashboards, tables, charts, theme modes and route-based workspaces.", "GUI connected to database."),
        ("Backend", "Next.js API route handlers", "Authenticated endpoints for CRUD, reports, auth, staff, settings and AI.", "Insert/update/delete/search APIs."),
        ("Database", "PostgreSQL", "Relational storage with PK/FK, enums, constraints and indexes.", "Database creation and integrity."),
        ("ORM", "Prisma Client", "Typed queries, relation includes, transactions and migrations.", "Front-end/database connectivity."),
        ("Validation", "Zod schemas", "Request parsing, required fields, email, number and enum validation.", "Debugged data entry forms."),
        ("AI", "Gemini API via @google/genai", "Business assistant with role-aware tools and approval-gated writes.", "Advanced reporting and tool workflow."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)
    add_heading(doc, "4. System Architecture", 1)
    doc.add_picture(str(arch_path), width=Cm(17.2))
    p(doc, "Figure 2. ShopIQ architecture from user roles to PostgreSQL and Gemini tools.", align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED, size=8.5)
    add_callout(doc, "Connectivity summary", "Every protected page reads the current user session, calls server-side data functions or API routes, and stores or retrieves records through Prisma. The route handlers return consistent JSON responses and use shared permission rules.", fill="F0FDFA", accent=TEAL)
    doc.add_page_break()


def add_roles(doc):
    add_heading(doc, "5. User Roles and Permission Design", 1)
    p(doc, "The sample OEL brief requires Admin and Normal User roles. ShopIQ expands this into a practical three-role model: Admin, Manager and Staff. Manager is useful in real shops because the owner may delegate operations without giving unrestricted control.")
    table = doc.add_table(rows=1, cols=5)
    headers = ["Resource", "Admin", "Manager", "Staff", "Design note"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    rows = [
        ("Products", "CRUD", "CRUD", "Read", "Stock master is protected from accidental staff edits."),
        ("Customers", "CRUD", "CRUD", "Create/Read/Update", "Staff can serve customers, but delete stays owner-level."),
        ("Suppliers", "CRUD", "CRUD", "No access", "Supplier cashflow is confidential."),
        ("Payments", "CRUD", "CRUD", "Create/Read customer receipts", "Supplier payouts are owner/manager only."),
        ("Invoices", "CRUD", "CRUD", "Create/Read", "Billing is available to daily staff."),
        ("Purchases", "CRUD", "CRUD", "No access", "Supplier procurement is owner/manager only."),
        ("Staff", "CRUD", "Manager can create/manage Staff", "No access", "Manager cannot promote themselves or create Admins."),
        ("Reports/Settings", "Read/Update", "Read/Update", "No access", "Strategic reports and settings remain protected."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)
    add_callout(doc, "Implementation evidence", "The permission matrix is implemented in src/lib/permissions.ts and then applied in API routes and UI pages. This means buttons are hidden/disabled in the interface and server routes still reject unauthorized requests.", fill="FFF7ED", accent=AMBER)
    doc.add_page_break()


def add_database_design(doc, erd_path):
    add_heading(doc, "6. Database Design and ERD", 1)
    p(doc, "The database is normalized around a Shop workspace. Almost every operational table includes shopId so records are isolated by shop. Transaction tables such as Invoice, Purchase and Payment separate header-level totals from line-level item rows, which avoids repeated data and supports accurate reporting.")
    doc.add_picture(str(erd_path), width=Cm(17.2))
    p(doc, "Figure 3. Simplified ERD for the ShopIQ relational database.", align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED, size=8.5)
    add_heading(doc, "Design principles used", 2)
    principles = [
        "Normalization: line items are separated from invoice and purchase headers.",
        "Referential integrity: foreign keys link users, shops, products, customers, suppliers and transactions.",
        "Controlled vocabularies: enums are used for roles, invoice status, payment direction, methods and stock movement types.",
        "Query performance: indexes target common filters such as shopId, status, date, balance and product stock.",
        "Auditability: ActivityLog and StockMovement preserve operational history.",
    ]
    for item in principles:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(item)
    doc.add_page_break()


def add_schema_tables(doc):
    add_heading(doc, "7. Tables, Keys, Constraints and Indexes", 1)
    table = doc.add_table(rows=1, cols=4)
    headers = ["Table", "Purpose", "Main keys/constraints", "Important indexes"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    rows = [
        ("Shop", "Top-level shop/workspace.", "id PK.", "createdAt."),
        ("User", "Team account and authentication role.", "id PK, email unique, shopId FK.", "shopId+role, shopId+status."),
        ("Category", "Product grouping.", "id PK, shopId FK, unique(shopId,name).", "shopId+createdAt."),
        ("Product", "Inventory master and pricing.", "id PK, unique(shopId,sku), categoryId FK.", "name, status, category, stock/reorder."),
        ("Customer", "Customer ledger and credit balance.", "id PK, shopId FK.", "name, phone, balance."),
        ("Supplier", "Supplier ledger and reliability.", "id PK, shopId FK.", "name, phone, balance."),
        ("Invoice", "Sales transaction header.", "id PK, unique(shopId,invoiceNo), customer/user FK.", "status, date, customer."),
        ("InvoiceItem", "Products sold in an invoice.", "id PK, invoiceId FK, productId FK.", "invoiceId, productId."),
        ("Purchase", "Supplier purchase header.", "id PK, unique(shopId,purchaseNo), supplier/user FK.", "status, supplier, date."),
        ("PurchaseItem", "Products received in a purchase.", "id PK, purchaseId FK, productId FK.", "purchaseId, productId."),
        ("Payment", "Customer receipts and supplier payouts.", "id PK, optional invoice/purchase/customer/supplier FK.", "direction, paidAt, customer, supplier."),
        ("StockMovement", "Stock history with before/after quantity.", "id PK, product/user FK.", "product, type, movedAt."),
        ("AssistantThread/Message", "AI conversation and action metadata.", "thread/message PKs, thread FK.", "threadId+createdAt."),
        ("ActivityLog", "Human-readable audit trail.", "id PK, shopId/user FK.", "shopId+createdAt, shopId+type."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)
    add_heading(doc, "Integrity choices", 2)
    p(doc, "Important delete behaviors are intentionally mixed. Shop deletion cascades through its child records, while invoice creator references use Restrict to preserve accountability. Optional relations such as customerId or supplierId may use SetNull so a transaction can remain historically valid even if a related party record is later removed.")
    doc.add_page_break()


def add_crud_validation(doc):
    add_heading(doc, "8. CRUD Operations and Validation", 1)
    p(doc, "ShopIQ implements CRUD through reusable workspace screens and dedicated API routes. Each module combines front-end forms with server-side Zod validation and Prisma database calls. The UI is role-aware, but server routes remain the final security layer.")
    table = doc.add_table(rows=1, cols=6)
    headers = ["Module", "Create", "Read/Search", "Update", "Delete", "Validation highlights"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    rows = [
        ("Products", "Yes", "List + search", "Yes", "Archive/delete route", "SKU, prices, stock, reorder fields."),
        ("Customers", "Yes", "Ledger/search", "Yes", "Owner-level", "Name, contact, email, credit limit."),
        ("Suppliers", "Admin/Manager", "Supplier cockpit", "Yes", "Owner-level", "Reliability score, balance, contact fields."),
        ("Invoices", "Yes", "Billing list", "Status update", "Owner-level", "Items, quantities, totals, due amount."),
        ("Purchases", "Admin/Manager", "Purchase list", "Status update", "Owner-level", "Supplier items and received totals."),
        ("Payments", "Yes", "Cashflow timeline", "Owner-level", "Owner-level", "Direction, method, amount, linked party."),
        ("Staff", "Admin/Manager", "Team list", "Role/status update", "Owner-level", "Role rules prevent unsafe promotions."),
        ("Settings", "No create", "Shop profile", "Admin/Manager", "No", "Shop name, city, phone, address."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)
    add_callout(doc, "Exception handling", "The shared API response utilities return unauthorized, forbidden and validation errors consistently. Zod catches missing or malformed fields before Prisma attempts a database write.", fill="FEF2F2", accent=RED)
    doc.add_page_break()


def add_queries_reports(doc):
    add_heading(doc, "9. Query Implementation and Reports", 1)
    p(doc, "The application uses select queries with relation includes, aggregation in the dashboard data layer, filtering by status/date, and ranked lists for operational decision making. Reports are visible in the dashboard and reports modules rather than exported as isolated static files.")
    table = doc.add_table(rows=1, cols=4)
    headers = ["Report / query", "Data used", "Query behavior", "Business value"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    rows = [
        ("Daily/active-day sales", "Invoice", "Date filtering and sum(total).", "Shows current sales even when demo data has no literal current-day records."),
        ("Monthly revenue", "Invoice", "Month window or latest active 30 days.", "Improves dashboard accuracy for seeded historical data."),
        ("Low stock risk", "Product", "stockQty <= reorderLevel.", "Highlights reorder needs."),
        ("Customer dues", "Customer + Invoice + Payment", "Balance ranking and payment history.", "Supports collection decisions."),
        ("Supplier payables", "Supplier + Purchase + Payment", "Role-filtered supplier balances.", "Protects sensitive cashflow."),
        ("Fast/slow movers", "StockMovement + Product", "SALE movement aggregation.", "Identifies demand and dead stock."),
        ("Payment method mix", "Payment", "Group by method and sum amount.", "Cashflow and channel insight."),
        ("Activity feed", "ActivityLog", "Recent events by shop.", "Audit trail for actions."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)
    add_heading(doc, "Professional report generation", 2)
    p(doc, "The reports are presented as charts, KPI cards, ranked lists and timelines. This is more useful for a shop owner than a raw SQL printout because it immediately answers: what sold, what is risky, who owes money, and what needs action.")
    doc.add_page_break()


def add_frontend_ai_testing(doc, ai_path):
    add_heading(doc, "10. Front-End and Database Connectivity", 1)
    p(doc, "The interface is built as a modern workspace with separate admin and staff routes. CRUD pages reuse a common manager component for search, details, form validation and action buttons. Route handlers connect to PostgreSQL through Prisma and are protected by JWT cookie sessions.")
    table = doc.add_table(rows=1, cols=3)
    headers = ["Area", "Implementation", "Evidence in project"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    rows = [
        ("Authentication", "JWT session cookie with bcrypt password hashes.", "src/lib/session.ts, src/lib/auth.ts, login/signup routes."),
        ("Workspace UI", "Admin and staff shells, sidebar, topbar, dashboards and CRUD pages.", "src/app/admin/*, src/app/staff/*, src/components/workspace/*."),
        ("Database access", "Prisma singleton client and server-side data functions.", "src/lib/prisma.ts, src/lib/data.ts."),
        ("Validation", "Zod schemas for API bodies and AI action payloads.", "src/lib/validation.ts, API route schemas."),
        ("Theme support", "Light/dark, classic/liquid glass, and shadcn/tweakcn-style presets.", "theme provider, globals.css and toggles."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)
    doc.add_page_break()

    add_heading(doc, "11. AI Assistant Feature System", 1)
    p(
        doc,
        "ShopIQ includes a live Gemini assistant built for the project itself, not a mock chatbot. It is connected to ShopIQ through the /api/ai/chat route, @google/genai, the Prisma-backed data layer, and the role rules already used by the application. The assistant can read live operating context, run decision-support jobs, answer exact money/date questions, and prepare database actions through a preview-first approval flow.",
    )
    doc.add_picture(str(ai_path), width=Cm(17.2))
    p(doc, "Figure 4. Actual ShopIQ AI feature system implemented through role-aware Gemini tools.", align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED, size=8.5)

    add_heading(doc, "AI tools implemented in ShopIQ", 2)
    tool_table = doc.add_table(rows=1, cols=4)
    headers = ["Tool", "Purpose in ShopIQ", "Uses live project data from", "Permission behavior"]
    for i, h in enumerate(headers):
        tool_table.rows[0].cells[i].text = h
    rows = [
        ("get_dashboard_snapshot", "Reads metrics, low stock, dues, activity and chart-ready dashboard data.", "src/lib/data.ts and Prisma models.", "Role-filtered supplier/cashflow data."),
        ("search_business_records", "Searches products, customers, suppliers, invoices, purchases, payments and staff before selecting ids.", "Prisma queries over ShopIQ records.", "Rejects entities the role cannot read."),
        ("get_record_details", "Loads full details for a selected visible record.", "Product, customer, supplier, invoice, purchase, payment and staff tables.", "Uses the same resource permission matrix."),
        ("run_operating_job", "Runs read-only jobs: reorder plan, collections plan, cashflow risk, sales quality review and stock audit.", "Dashboard snapshot, products, customers, suppliers, payments and movements.", "Supplier-side jobs require Admin/Manager."),
        ("get_sales_summary", "Answers exact revenue, earnings, gross profit, cash received, item count and invoice totals for a date range.", "Invoice, InvoiceItem, Payment and Product data.", "Read-only and shop-scoped."),
        ("get_customer_balance_summary", "Answers pending money, outstanding invoices and recent payments for one customer.", "Customer, Invoice and Payment tables.", "Customer read permission required."),
        ("prepare_business_action", "Validates write requests and creates a preview. It never writes immediately.", "Zod schemas and actionResource permissions.", "Requires later user approval before execution."),
    ]
    for row in rows:
        cells = tool_table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(tool_table)

    doc.add_page_break()
    add_heading(doc, "AI write actions supported", 2)
    action_table = doc.add_table(rows=1, cols=4)
    headers = ["Action family", "Supported actions", "Validation source", "Execution result"]
    for i, h in enumerate(headers):
        action_table.rows[0].cells[i].text = h
    rows = [
        ("Inventory", "create_product, update_product", "Product Zod schemas: name, SKU/category, prices, stock and reorder fields.", "Creates/updates product records and logs AI product activity."),
        ("Customer ledger", "create_customer, update_customer", "Customer schemas: name, contact, email, creditLimit, balance and notes.", "Creates/updates customer records and opens Customers workspace."),
        ("Supplier ledger", "create_supplier, update_supplier", "Supplier schemas: contact, balance, reliabilityScore and notes.", "Admin/Manager only because supplier cashflow is restricted."),
        ("Cashflow", "create_payment", "Payment schema: direction, method, amount, party and linked invoice/purchase.", "Customer receipts for allowed users; supplier payouts only for Admin/Manager."),
        ("Sales and procurement", "create_invoice, create_purchase", "Item schemas require product id/SKU/name, quantities and prices/costs.", "Creates transaction headers/items and updates related operational data."),
        ("Team management", "create_staff, update_staff", "Staff schemas plus canCreateStaffRole/canManageStaffMember rules.", "Admin can manage all roles; Manager can create/manage Staff only."),
    ]
    for row in rows:
        cells = action_table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(action_table)

    add_heading(doc, "AI guardrails and user experience", 2)
    guardrails = [
        "No mock fallback: if GEMINI_API_KEY is missing, the API returns a configuration error instead of fake answers.",
        "Current chat context is sent to Gemini so follow-up prompts are understood within the same conversation.",
        "Every write task is confirmation-gated: the assistant creates a preview first, and the server executes only after approval.",
        "Pending actions are stored in AssistantMessage metadata with previewId and status so stale approvals cannot accidentally execute another action.",
        "The system prompt blocks deletes, account suspension, bulk destructive edits, secret retrieval and role bypassing.",
        "The assistant is instructed to choose ids only from tool results and never invent SKUs, invoice numbers or balances.",
        "The AI tab is focused on the assistant only, while charts remain in dashboard/report modules.",
    ]
    for item in guardrails:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(item)

    add_callout(
        doc,
        "Project-specific AI conclusion",
        "The assistant is a database-aware operating layer for ShopIQ. Its feature set is limited to the modules that exist in the project: inventory, customers, suppliers, invoices, purchases, payments, staff, reports and dashboard context.",
        fill="F5F3FF",
        accent="8B5CF6",
    )
    doc.add_page_break()

    add_heading(doc, "12. Testing, Exception Handling and Rubric Mapping", 1)
    table2 = doc.add_table(rows=1, cols=3)
    headers = ["Rubric criterion", "How ShopIQ satisfies it", "Evidence"]
    for i, h in enumerate(headers):
        table2.rows[0].cells[i].text = h
    rows = [
        ("Database creation", "Tables, enums, relationships and migrations are defined.", "Prisma schema and SQL migration."),
        ("Keys and constraints", "PK, FK, unique constraints, onDelete rules and indexes are present.", "schema.prisma relation annotations."),
        ("Insert/update/delete", "CRUD routes exist for operational modules.", "API route handlers and workspace CRUD UI."),
        ("Queries", "Filtering, relation includes, aggregation and ranking are used.", "src/lib/data.ts and reports routes."),
        ("Connectivity", "Next.js front end calls API routes backed by Prisma/PostgreSQL.", "Route handlers + Prisma client."),
        ("Reports", "Dashboards and reports produce operational metrics and charts.", "Dashboard snapshot and report pages."),
        ("Understanding", "Role model, AI guardrails, validation and exception handling are documented.", "This report and implemented source files."),
    ]
    for row in rows:
        cells = table2.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table2)
    doc.add_page_break()


def add_appendix(doc):
    add_heading(doc, "Appendix A. Representative SQL / Prisma Queries", 1)
    examples = [
        ("Low-stock products", "SELECT name, sku, stockQty, reorderLevel FROM Product WHERE shopId = ? AND status = 'ACTIVE' AND stockQty <= reorderLevel ORDER BY stockQty ASC;"),
        ("Customer due ranking", "SELECT name, phone, balance FROM Customer WHERE shopId = ? AND balance > 0 ORDER BY balance DESC LIMIT 10;"),
        ("Invoice status report", "SELECT status, COUNT(*) AS count, SUM(total) AS gross FROM Invoice WHERE shopId = ? GROUP BY status;"),
        ("Daily revenue", "SELECT DATE(invoiceDate) AS day, SUM(total) AS revenue FROM Invoice WHERE shopId = ? GROUP BY DATE(invoiceDate) ORDER BY day;"),
        ("Product movement history", "SELECT p.name, m.type, m.quantity, m.beforeQty, m.afterQty, m.movedAt FROM StockMovement m JOIN Product p ON p.id = m.productId WHERE m.shopId = ? ORDER BY m.movedAt DESC;"),
    ]
    for title, sql in examples:
        para = doc.add_paragraph()
        r = para.add_run(title)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(TEAL)
        code = doc.add_paragraph()
        code.paragraph_format.left_indent = Cm(0.4)
        code.paragraph_format.space_after = Pt(8)
        rr = code.add_run(sql)
        rr.font.name = "Consolas"
        rr.font.size = Pt(8.8)
        rr.font.color.rgb = RGBColor.from_string(INK)

    add_heading(doc, "Conclusion", 1)
    p(doc, "ShopIQ fulfills the OEL requirement as a complete database application: it creates a relational schema, enforces keys and constraints, implements role-based CRUD operations, connects a polished front end to PostgreSQL, and generates useful operational reports. The project goes beyond a basic DBMS submission by adding a safe AI assistant, multi-role workflows, audit logging, and a modern UI system while keeping the database at the center of the application.")


def add_metadata(doc):
    props = doc.core_properties
    props.title = "ShopIQ OEL DBMS Report"
    props.subject = "CS 222 Database Management Systems Open Ended Lab"
    props.keywords = "ShopIQ, DBMS, PostgreSQL, Prisma, Next.js, OEL"
    props.comments = "Generated project report based on OEL-DBMS-2026 sample brief and ShopIQ source code."


def build_doc():
    ensure_dirs()
    arch = architecture_image()
    erd = erd_image()
    rubric = rubric_image()
    ai_features = ai_features_image()
    logo = shopiq_wordmark()

    doc = Document()
    setup_sections(doc)
    set_doc_styles(doc)
    add_metadata(doc)
    add_cover(doc, logo)
    start_body_section(doc)
    add_contents(doc)
    add_project_overview(doc, rubric)
    add_problem_objectives(doc)
    add_stack_architecture(doc, arch)
    add_roles(doc)
    add_database_design(doc, erd)
    add_schema_tables(doc)
    add_crud_validation(doc)
    add_queries_reports(doc)
    add_frontend_ai_testing(doc, ai_features)
    add_appendix(doc)
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    build_doc()
